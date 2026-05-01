from docx import Document

def save_summary_to_word(summary, output_file="output.docx"):
    doc = Document()

    doc.add_heading("AI Generated Medical Summary", level=1)

    doc.add_paragraph(summary)

    doc.save(output_file)

    print(f"Word file saved as: {output_file}")