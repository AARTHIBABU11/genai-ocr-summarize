from summarizer import summarize_long_text
from ocr import extract_text
from docx import Document   # 🔥 NEW IMPORT

# 1. File path
file_path = "sample.pdf"

# 2. OCR step
text = extract_text(file_path)

# 3. NLP summarization
summary = summarize_long_text(text)

# 4. Output in terminal
print("\n===== SUMMARY =====\n")
print(summary)

# ----------------------------
# 5. SAVE TO WORD FILE (.docx)
# ----------------------------
doc = Document()

doc.add_heading("AI Generated Summary Report", level=1)

doc.add_paragraph(summary)

output_file = "output.docx"
doc.save(output_file)

print(f"\n✅ Word file created: {output_file}")